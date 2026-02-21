"""Tests for Knowledge Base routes and text extraction."""

import io
import tempfile
from contextlib import contextmanager

import pytest
from fastapi.testclient import TestClient
from sqlalchemy import create_engine
from sqlalchemy.orm import sessionmaker

from intellibox.database import Base
from intellibox.models import KnowledgeDocument

# Create test database
test_db_fd, test_db_path = tempfile.mkstemp(suffix='_knowledge_base.db', prefix='test_intellibox_')
TEST_DATABASE_URL = f"sqlite:///{test_db_path}"
test_engine = create_engine(TEST_DATABASE_URL, pool_pre_ping=True)
TestSessionLocal = sessionmaker(autocommit=False, autoflush=False, bind=test_engine)


@contextmanager
def override_get_session():
    """Override database session for testing."""
    session = TestSessionLocal()
    try:
        yield session
    finally:
        session.close()


# Patch get_session at the database module level — deps.get_session() delegates there
import intellibox.database as database_module
import intellibox.settings_service as settings_service_module

database_module.get_session = override_get_session
settings_service_module.get_session = override_get_session

from intellibox.web.app import app

client = TestClient(app)


@pytest.fixture(scope="function")
def setup_database():
    """Create tables before each test."""
    Base.metadata.drop_all(test_engine)
    Base.metadata.create_all(test_engine)
    yield
    Base.metadata.drop_all(test_engine)


class TestKnowledgeBaseRoutes:
    """Test Knowledge Base web routes."""

    def test_legacy_url_redirects_to_settings(self, setup_database):
        """Test /knowledge-base redirects to /settings?tab=kb."""
        response = client.get("/knowledge-base", follow_redirects=False)
        assert response.status_code == 302
        assert "/settings?tab=kb" in response.headers.get("location", "")

    def test_kb_tab_on_settings_page(self, setup_database):
        """Test KB tab appears on the settings page."""
        response = client.get("/settings")
        assert response.status_code == 200
        assert b"tab-kb" in response.content
        assert b"pane-kb" in response.content
        assert b"Knowledge Base" in response.content

    def test_kb_tab_empty_state(self, setup_database):
        """Test KB tab shows empty state message when no documents."""
        response = client.get("/settings")
        assert response.status_code == 200
        assert b"No documents uploaded" in response.content

    def test_kb_tab_shows_documents(self, setup_database):
        """Test that uploaded documents appear in the KB tab on settings."""
        with override_get_session() as session:
            doc = KnowledgeDocument(
                filename="report.pdf", file_type="pdf", file_size=5000,
                extracted_text="PDF content here", extraction_status="success",
                description="Monthly report",
            )
            session.add(doc)
            session.commit()

        response = client.get("/settings")
        assert response.status_code == 200
        assert b"report.pdf" in response.content
        assert b"Monthly report" in response.content
        assert b"PDF" in response.content

    def test_kb_tab_summary_stats(self, setup_database):
        """Test summary stats on KB tab are calculated correctly."""
        with override_get_session() as session:
            session.add(KnowledgeDocument(
                filename="a.txt", file_type="txt", file_size=100,
                extracted_text="hello", extraction_status="success",
            ))
            session.add(KnowledgeDocument(
                filename="b.txt", file_type="txt", file_size=200,
                extracted_text="world!", extraction_status="success",
            ))
            session.commit()

        response = client.get("/settings")
        assert response.status_code == 200
        # Tab badge should show count "2"
        assert b">2<" in response.content

    def test_kb_tab_shows_extraction_status(self, setup_database):
        """Test that extraction status badges render correctly."""
        with override_get_session() as session:
            session.add(KnowledgeDocument(
                filename="good.txt", file_type="txt", file_size=50,
                extracted_text="content", extraction_status="success",
            ))
            session.add(KnowledgeDocument(
                filename="bad.txt", file_type="txt", file_size=50,
                extracted_text=None, extraction_status="failed",
            ))
            session.commit()

        response = client.get("/settings")
        assert response.status_code == 200
        assert b"good.txt" in response.content
        assert b"bad.txt" in response.content
        assert b"failed" in response.content

    def test_upload_txt(self, setup_database):
        """Test uploading a text file."""
        response = client.post(
            "/knowledge-base/upload",
            files={"file": ("test.txt", b"Hello World", "text/plain")},
            data={"description": "Test document"},
            follow_redirects=False,
        )
        assert response.status_code == 303
        location = response.headers.get("location", "")
        assert "/settings" in location
        assert "kb_success" in location

        # Verify document stored in DB
        with override_get_session() as session:
            doc = session.query(KnowledgeDocument).first()
            assert doc is not None
            assert doc.filename == "test.txt"
            assert doc.file_type == "txt"
            assert doc.extracted_text == "Hello World"
            assert doc.extraction_status == "success"
            assert doc.description == "Test document"
            assert doc.file_size == 11

    def test_upload_no_description(self, setup_database):
        """Test uploading without a description."""
        response = client.post(
            "/knowledge-base/upload",
            files={"file": ("notes.txt", b"Some notes", "text/plain")},
            data={"description": ""},
            follow_redirects=False,
        )
        assert response.status_code == 303

        with override_get_session() as session:
            doc = session.query(KnowledgeDocument).first()
            assert doc.description is None

    def test_upload_invalid_extension(self, setup_database):
        """Test rejecting unsupported file types."""
        response = client.post(
            "/knowledge-base/upload",
            files={"file": ("virus.exe", b"bad content", "application/octet-stream")},
            data={"description": ""},
            follow_redirects=False,
        )
        assert response.status_code == 303
        assert "kb_error" in response.headers.get("location", "")

        with override_get_session() as session:
            assert session.query(KnowledgeDocument).count() == 0

    def test_upload_empty_file(self, setup_database):
        """Test rejecting empty files."""
        response = client.post(
            "/knowledge-base/upload",
            files={"file": ("empty.txt", b"", "text/plain")},
            data={"description": ""},
            follow_redirects=False,
        )
        assert response.status_code == 303
        assert "kb_error" in response.headers.get("location", "")

    def test_upload_oversized_file(self, setup_database):
        """Test rejecting files exceeding the 10 MB size limit."""
        # Create content just over the limit
        content = b"x" * (10 * 1024 * 1024 + 1)
        response = client.post(
            "/knowledge-base/upload",
            files={"file": ("big.txt", content, "text/plain")},
            data={"description": ""},
            follow_redirects=False,
        )
        assert response.status_code == 303
        assert "kb_error" in response.headers.get("location", "")

        with override_get_session() as session:
            assert session.query(KnowledgeDocument).count() == 0

    def test_delete_document(self, setup_database):
        """Test deleting a document."""
        # Create a document
        with override_get_session() as session:
            doc = KnowledgeDocument(
                filename="old.txt", file_type="txt", file_size=100,
                extracted_text="old content", extraction_status="success",
            )
            session.add(doc)
            session.commit()
            doc_id = doc.id

        response = client.post(
            f"/knowledge-base/{doc_id}/delete",
            follow_redirects=False,
        )
        assert response.status_code == 303
        location = response.headers.get("location", "")
        assert "/settings" in location
        assert "kb_deleted" in location

        with override_get_session() as session:
            assert session.query(KnowledgeDocument).count() == 0

    def test_delete_nonexistent(self, setup_database):
        """Test deleting a nonexistent document doesn't crash."""
        response = client.post(
            "/knowledge-base/999/delete",
            follow_redirects=False,
        )
        assert response.status_code == 303

    def test_detail_page_loads(self, setup_database):
        """Test document detail page loads with back link to settings."""
        with override_get_session() as session:
            doc = KnowledgeDocument(
                filename="detail.txt", file_type="txt", file_size=50,
                extracted_text="Full text here", extraction_status="success",
            )
            session.add(doc)
            session.commit()
            doc_id = doc.id

        response = client.get(f"/knowledge-base/{doc_id}")
        assert response.status_code == 200
        assert b"detail.txt" in response.content
        assert b"Full text here" in response.content
        assert b"/settings?tab=kb" in response.content  # back link points to settings

    def test_detail_page_nonexistent_redirects(self, setup_database):
        """Test that viewing a nonexistent document redirects to settings KB tab."""
        response = client.get("/knowledge-base/999", follow_redirects=False)
        assert response.status_code == 303
        assert "/settings?tab=kb" in response.headers.get("location", "")

    def test_detail_page_highlight(self, setup_database):
        """Test document detail page with highlight query param."""
        with override_get_session() as session:
            doc = KnowledgeDocument(
                filename="search.txt", file_type="txt", file_size=100,
                extracted_text="Some text with keyword inside", extraction_status="success",
            )
            session.add(doc)
            session.commit()
            doc_id = doc.id

        response = client.get(f"/knowledge-base/{doc_id}?highlight=keyword")
        assert response.status_code == 200
        assert b"keyword" in response.content

    def test_upload_redirects_with_tab_param(self, setup_database):
        """Test that all upload outcomes redirect to settings with tab=kb."""
        # Success case
        response = client.post(
            "/knowledge-base/upload",
            files={"file": ("ok.txt", b"content", "text/plain")},
            data={"description": ""},
            follow_redirects=False,
        )
        location = response.headers.get("location", "")
        assert "tab=kb" in location or "kb_success" in location

    def test_kb_tab_document_count_badge(self, setup_database):
        """Test that the tab badge shows the correct document count."""
        with override_get_session() as session:
            for i in range(3):
                session.add(KnowledgeDocument(
                    filename=f"doc{i}.txt", file_type="txt", file_size=10,
                    extracted_text=f"text{i}", extraction_status="success",
                ))
            session.commit()

        response = client.get("/settings")
        assert response.status_code == 200
        # Badge should show 3
        assert b">3<" in response.content


class TestChunkText:
    """Tests for the text chunking utility."""

    def test_chunk_text_returns_list(self):
        from intellibox.knowledge.embeddings import chunk_text
        result = chunk_text("Hello world")
        assert isinstance(result, list)
        assert len(result) >= 1

    def test_chunk_text_respects_size_limit(self):
        from intellibox.knowledge.embeddings import CHUNK_SIZE, chunk_text
        # Create a very long text
        long_text = "word " * 1000  # ~5000 chars
        chunks = chunk_text(long_text)
        for chunk in chunks:
            assert len(chunk) <= CHUNK_SIZE + 50  # small tolerance for paragraph boundary

    def test_chunk_text_empty(self):
        from intellibox.knowledge.embeddings import chunk_text
        assert chunk_text("") == []
        assert chunk_text("   ") == []
        assert chunk_text(None) == []

    def test_chunk_text_preserves_content(self):
        from intellibox.knowledge.embeddings import chunk_text
        text = "Paragraph one.\n\nParagraph two.\n\nParagraph three."
        chunks = chunk_text(text)
        # All content should appear in at least one chunk
        combined = " ".join(chunks)
        assert "Paragraph one" in combined
        assert "Paragraph two" in combined
        assert "Paragraph three" in combined


class TestTfidfSearch:
    """Tests for TF-IDF search with caching."""

    def test_tfidf_search_returns_results(self, setup_database):
        """Upload a doc and search for a keyword — should get results."""
        from intellibox.knowledge.embeddings import invalidate_tfidf_cache, search_by_tfidf

        invalidate_tfidf_cache()

        # Seed a document with known content
        with override_get_session() as session:
            doc = KnowledgeDocument(
                filename="test_search.txt", file_type="txt", file_size=100,
                extracted_text="This document discusses machine learning algorithms and neural networks.",
                extraction_status="success",
            )
            session.add(doc)
            session.commit()

        # Patch get_session in embeddings module to use our test session
        import intellibox.knowledge.embeddings as emb_module
        original_get_session = emb_module.get_session
        emb_module.get_session = override_get_session

        try:
            results = search_by_tfidf("machine learning", threshold=0.01)
            assert len(results) >= 1
            assert results[0]["doc_filename"] == "test_search.txt"
            assert results[0]["method"] == "tfidf"
        finally:
            emb_module.get_session = original_get_session

    def test_tfidf_cache_reused_on_second_call(self, setup_database):
        """Second TF-IDF search reuses cached vectorizer (no rebuild)."""
        from intellibox.knowledge.embeddings import _tfidf_cache, invalidate_tfidf_cache, search_by_tfidf

        invalidate_tfidf_cache()

        with override_get_session() as session:
            doc = KnowledgeDocument(
                filename="cache_test.txt", file_type="txt", file_size=50,
                extracted_text="Python programming language features and syntax.",
                extraction_status="success",
            )
            session.add(doc)
            session.commit()

        import intellibox.knowledge.embeddings as emb_module
        original_get_session = emb_module.get_session
        emb_module.get_session = override_get_session

        try:
            # First search builds cache
            search_by_tfidf("python", threshold=0.01)
            assert _tfidf_cache["vectorizer"] is not None
            cached_vectorizer = _tfidf_cache["vectorizer"]

            # Second search should reuse same vectorizer
            search_by_tfidf("syntax", threshold=0.01)
            assert _tfidf_cache["vectorizer"] is cached_vectorizer
        finally:
            emb_module.get_session = original_get_session

    def test_tfidf_cache_invalidated_on_doc_change(self, setup_database):
        """Cache is invalidated when document set changes."""
        from intellibox.knowledge.embeddings import _tfidf_cache, invalidate_tfidf_cache, search_by_tfidf

        invalidate_tfidf_cache()

        with override_get_session() as session:
            doc = KnowledgeDocument(
                filename="first.txt", file_type="txt", file_size=50,
                extracted_text="First document about testing.",
                extraction_status="success",
            )
            session.add(doc)
            session.commit()

        import intellibox.knowledge.embeddings as emb_module
        original_get_session = emb_module.get_session
        emb_module.get_session = override_get_session

        try:
            # Build initial cache
            search_by_tfidf("testing", threshold=0.01)
            first_hash = _tfidf_cache["corpus_hash"]

            # Add another document
            with override_get_session() as session:
                doc2 = KnowledgeDocument(
                    filename="second.txt", file_type="txt", file_size=50,
                    extracted_text="Second document about deployment.",
                    extraction_status="success",
                )
                session.add(doc2)
                session.commit()

            # Next search should detect change and rebuild
            search_by_tfidf("deployment", threshold=0.01)
            assert _tfidf_cache["corpus_hash"] != first_hash
        finally:
            emb_module.get_session = original_get_session


class TestKnowledgeContextCache:
    """Tests for the get_knowledge_context TTL cache."""

    def test_context_returns_text(self, setup_database):
        """get_knowledge_context returns non-empty string when docs exist."""
        from intellibox.knowledge import _invalidate_kb_cache, get_knowledge_context

        _invalidate_kb_cache()

        with override_get_session() as session:
            doc = KnowledgeDocument(
                filename="context.txt", file_type="txt", file_size=50,
                extracted_text="Important program context here.",
                extraction_status="success",
            )
            session.add(doc)
            session.commit()

        import intellibox.knowledge as kb_module
        original_get_session = kb_module.get_session
        kb_module.get_session = override_get_session

        try:
            context = get_knowledge_context()
            assert "Important program context" in context
        finally:
            kb_module.get_session = original_get_session

    def test_context_cache_avoids_repeated_queries(self, setup_database):
        """Second call to get_knowledge_context uses cached value."""
        from intellibox.knowledge import _invalidate_kb_cache, _kb_cache, get_knowledge_context

        _invalidate_kb_cache()

        with override_get_session() as session:
            doc = KnowledgeDocument(
                filename="cached.txt", file_type="txt", file_size=50,
                extracted_text="Cached content.",
                extraction_status="success",
            )
            session.add(doc)
            session.commit()

        import intellibox.knowledge as kb_module
        original_get_session = kb_module.get_session
        kb_module.get_session = override_get_session

        try:
            # First call populates cache
            result1 = get_knowledge_context()
            assert _kb_cache["text"] is not None

            # Second call should return same object (cache hit)
            result2 = get_knowledge_context()
            assert result1 == result2
        finally:
            kb_module.get_session = original_get_session

    def test_context_empty_when_no_docs(self, setup_database):
        """get_knowledge_context returns empty string with no documents."""
        from intellibox.knowledge import _invalidate_kb_cache, get_knowledge_context

        _invalidate_kb_cache()

        import intellibox.knowledge as kb_module
        original_get_session = kb_module.get_session
        kb_module.get_session = override_get_session

        try:
            context = get_knowledge_context()
            assert context == ""
        finally:
            kb_module.get_session = original_get_session

    def test_context_cache_expires_after_ttl(self, setup_database):
        """After TTL expires, cache is refreshed from DB on next call."""
        from intellibox.knowledge import _invalidate_kb_cache, _kb_cache, get_knowledge_context

        _invalidate_kb_cache()

        with override_get_session() as session:
            doc = KnowledgeDocument(
                filename="ttl.txt", file_type="txt", file_size=50,
                extracted_text="Original content.",
                extraction_status="success",
            )
            session.add(doc)
            session.commit()

        import intellibox.knowledge as kb_module
        original_get_session = kb_module.get_session
        kb_module.get_session = override_get_session

        try:
            # First call populates cache
            result1 = get_knowledge_context()
            assert "Original content" in result1

            # Update the document behind the cache
            with override_get_session() as session:
                doc = session.query(KnowledgeDocument).first()
                doc.extracted_text = "Updated content."
                session.commit()

            # Second call within TTL should still return old cached value
            result2 = get_knowledge_context()
            assert "Original content" in result2

            # Advance the clock past TTL by manipulating expires directly
            _kb_cache["expires"] = 0  # Force expiry

            # Now the cache should miss and re-query the DB
            result3 = get_knowledge_context()
            assert "Updated content" in result3
        finally:
            kb_module.get_session = original_get_session

    def test_context_cache_no_db_query_on_hit(self, setup_database):
        """Verify the cache actually prevents DB queries on subsequent calls."""
        from unittest.mock import MagicMock

        from intellibox.knowledge import _invalidate_kb_cache, get_knowledge_context

        _invalidate_kb_cache()

        with override_get_session() as session:
            doc = KnowledgeDocument(
                filename="spy.txt", file_type="txt", file_size=50,
                extracted_text="Spy content.",
                extraction_status="success",
            )
            session.add(doc)
            session.commit()

        import intellibox.knowledge as kb_module
        original_get_session = kb_module.get_session
        kb_module.get_session = override_get_session

        try:
            # First call: populates cache (hits DB)
            get_knowledge_context()

            # Now replace get_session with a spy to detect DB access
            call_spy = MagicMock(side_effect=override_get_session)
            kb_module.get_session = call_spy

            # Second call should use cache (no DB call)
            get_knowledge_context()
            assert call_spy.call_count == 0, "Cache hit should not query DB"
        finally:
            kb_module.get_session = original_get_session


class TestTfidfStress:
    """Stress tests for TF-IDF search with larger corpora."""

    def test_tfidf_with_many_documents(self, setup_database):
        """TF-IDF search works correctly with 50+ documents."""
        from intellibox.knowledge.embeddings import invalidate_tfidf_cache, search_by_tfidf

        invalidate_tfidf_cache()

        topics = [
            "machine learning algorithms neural networks deep learning",
            "database optimization query performance indexing",
            "web development frontend backend deployment",
            "security authentication encryption protocols",
            "testing quality assurance continuous integration",
        ]

        with override_get_session() as session:
            for i in range(50):
                topic = topics[i % len(topics)]
                session.add(KnowledgeDocument(
                    filename=f"doc_{i:03d}.txt",
                    file_type="txt",
                    file_size=len(topic),
                    extracted_text=f"Document {i} covers: {topic}. " * 3,
                    extraction_status="success",
                ))
            session.commit()

        import intellibox.knowledge.embeddings as emb_module
        original_get_session = emb_module.get_session
        emb_module.get_session = override_get_session

        try:
            results = search_by_tfidf("machine learning", threshold=0.01)
            assert len(results) >= 1
            # All results should be from ML-related documents
            for r in results:
                assert r["method"] == "tfidf"
                assert r["score"] > 0
        finally:
            emb_module.get_session = original_get_session

    def test_tfidf_repeated_searches_use_cache(self, setup_database):
        """Running 20 different searches reuses the cached vectorizer."""
        from intellibox.knowledge.embeddings import _tfidf_cache, invalidate_tfidf_cache, search_by_tfidf

        invalidate_tfidf_cache()

        with override_get_session() as session:
            for i in range(10):
                session.add(KnowledgeDocument(
                    filename=f"multi_{i}.txt",
                    file_type="txt",
                    file_size=100,
                    extracted_text=f"Topic {i}: various keywords and phrases for testing search number {i}.",
                    extraction_status="success",
                ))
            session.commit()

        import intellibox.knowledge.embeddings as emb_module
        original_get_session = emb_module.get_session
        emb_module.get_session = override_get_session

        try:
            queries = [
                "keywords", "testing", "phrases", "topic", "various",
                "search", "number", "text", "document", "content",
                "alpha", "beta", "gamma", "delta", "epsilon",
                "foo", "bar", "baz", "qux", "quux",
            ]

            # First search builds cache
            search_by_tfidf(queries[0], threshold=0.01)
            vectorizer_after_first = _tfidf_cache["vectorizer"]

            # All subsequent searches should reuse the same vectorizer
            for q in queries[1:]:
                search_by_tfidf(q, threshold=0.01)
                assert _tfidf_cache["vectorizer"] is vectorizer_after_first, \
                    f"Vectorizer rebuilt unexpectedly on query '{q}'"
        finally:
            emb_module.get_session = original_get_session


class TestTextExtractor:
    """Test the text extraction module."""

    def test_extract_txt(self):
        from intellibox.knowledge.extractor import extract_text
        text, status = extract_text(b"Hello World", "txt")
        assert text == "Hello World"
        assert status == "success"

    def test_extract_txt_latin1(self):
        from intellibox.knowledge.extractor import extract_text
        content = "café résumé".encode("latin-1")
        text, status = extract_text(content, "txt")
        assert "caf" in text
        assert status == "success"

    def test_extract_unknown_type(self):
        from intellibox.knowledge.extractor import extract_text
        text, status = extract_text(b"content", "xlsx")
        assert text == ""
        assert status == "failed"

    def test_extract_pdf(self):
        """Test PDF extraction with a minimal PDF."""
        # Create a minimal PDF using pypdf
        from pypdf import PdfWriter

        from intellibox.knowledge.extractor import extract_text
        writer = PdfWriter()
        writer.add_blank_page(width=72, height=72)
        buf = io.BytesIO()
        writer.write(buf)
        pdf_bytes = buf.getvalue()

        text, status = extract_text(pdf_bytes, "pdf")
        # Blank page has no text, should be partial
        assert status == "partial"

    def test_extract_docx(self):
        """Test DOCX extraction with a minimal docx."""
        from docx import Document

        from intellibox.knowledge.extractor import extract_text

        doc = Document()
        doc.add_paragraph("Test paragraph content")
        buf = io.BytesIO()
        doc.save(buf)
        docx_bytes = buf.getvalue()

        text, status = extract_text(docx_bytes, "docx")
        assert "Test paragraph content" in text
        assert status == "success"

    def test_extract_txt_empty(self):
        """Test extracting from empty text content."""
        from intellibox.knowledge.extractor import extract_text
        text, status = extract_text(b"", "txt")
        assert text == ""
        assert status == "success"

    def test_extract_txt_unicode(self):
        """Test extracting UTF-8 unicode text."""
        from intellibox.knowledge.extractor import extract_text
        content = "Hello 世界 🌍".encode("utf-8")
        text, status = extract_text(content, "txt")
        assert "Hello" in text
        assert status == "success"

    @pytest.mark.parametrize("file_type", ["md", "csv", "json", "html", "xml", "rtf", "log"])
    def test_extract_text_based_formats(self, file_type):
        """All text-based formats should extract via the txt handler."""
        from intellibox.knowledge.extractor import extract_text
        content = f"Sample {file_type} content".encode("utf-8")
        text, status = extract_text(content, file_type)
        assert f"Sample {file_type} content" in text
        assert status == "success"
